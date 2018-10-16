from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches

#打开源xlsx文档和目标docx文档
document = Document()
wb = load_workbook('RG2018.xlsx')

#测试用例正文起始行
first_line = 5

#测试用例章节序号
title_start_id = 3

#生成测试用例表格
def write_doc_table(document, contents):
    #表头内容
    table_heads = ['模块', '功能', '测试项', '测试拓扑', '测试步骤', '预期结果', '测试结果', '备注']
    #添加表格
    table = document.add_table(rows=8, cols=0, style='Table Grid')
    #新增表头列col1，宽度1.2
    col1 = table.add_column(Inches(1.2))
    #新增内容列col2，宽度4.8
    col2 = table.add_column(Inches(4.8))
    #迭代填写单元格内容
    for cell_id,table_head,content in zip(range(8), table_heads, contents):
        col1.cells[cell_id].text = table_head
        #cells.cesll.text内容类型不能为None，因此转换为str类型
        col2.cells[cell_id].text = str(content)

#迭代sheet表，序号1-18，序号10为应用代理，功能暂时关闭
for i in [ i for i in range(1,3) if i != 10]:
    #设置当前活动sheet
    current_sheet = wb.sheetnames[i]
    ws = wb[current_sheet]

    #初始化标题序号
    module_seq = 0
    function_seq = 0
    case_seq = 0

    current_module = None
    current_function = None

    title_l2_seq = str(title_start_id) +'.' + str(i)
    title_l2 = title_l2_seq + ' ' + current_sheet
    document.add_heading(title_l2, level=2)

    #从测试用例内容起始行开始逐行迭代
    for i in range (first_line, ws.max_row):
        cid = ws[i][0].value
        module = ws[i][1].value
        function = ws[i][2].value
        case = ws[i][3].value
        level = ws[i][4].value
        steps = ws[i][5].value
        presult = ws[i][6].value
        result = ws[i][7].value
        beizhu = ws[i][8].value

        #获取的模块名为空时，沿用上一次获取的模块名
        if module == None:
            module = current_module
        #获取的模块名非空，说明新模块的开始，模块序号+1，同时重置功能序号，并拼接标题序号和内容
        else:
            module_seq = module_seq+1
            function_seq = 0
            title_l3_seq = title_l2_seq + '.' + str(module_seq) 
            title_l3 = title_l3_seq + ' ' + module
            document.add_heading(title_l3, level=3)
            current_module = module
        
        #获取的功能名为空时，沿用上一次获取的功能名
        if function == None:
            function = current_function
        #获取的功能名非空时，说明新功能块的开始，功能序号+1，同时重置测试项序号，并拼接标题序号和内容
        else:
            function_seq = function_seq + 1
            case_seq = 0
            title_l4_seq = title_l3_seq + '.' + str(function_seq)
            title_l4 = title_l4_seq + ' ' + function
            document.add_heading(title_l4, level=4)
            current_function = function
        
        #生成传递给write_doc_table函数需要的参数
        records = [module, function, case, ' ', steps, presult, result, beizhu]
        #测试项序号+1，同时拼接标题序号和内容
        case_seq = case_seq + 1
        title_l5_seq = title_l4_seq + '.' + str(case_seq)
        title_l5 = title_l5_seq + ' ' + case
        document.add_heading(title_l5, level=5)
        #生成该测试项的表格
        write_doc_table(document, records)
#保存目标docx文档
out_dir = 'output/'
out_filename = out_dir + 'RG2018测试报告.docx'
document.save(out_filename)