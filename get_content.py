from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches
import datetime
from docxtpl import DocxTemplate,RichText

document = Document()
wb = load_workbook('RG2018.xlsx')
first_line = 5

def write_doc_table(document, contents):
    table_heads = ['模块', '功能', '测试项', '测试拓扑', '测试步骤', '预期结果', '测试结果', '备注']
    table = document.add_table(rows=8, cols=0, style='Table Grid')
    col1 = table.add_column(Inches(1.2))
    col2 = table.add_column(Inches(4.8))
    for cell_id,table_head,content in zip(range(8), table_heads, contents):
        col1.cells[cell_id].text = table_head
        col2.cells[cell_id].text = str(content)

for i in [ i for i in range(1,17) if i != 10]:
    current_sheet = wb.sheetnames[i]
    ws = wb[current_sheet]

    current_module = None
    current_function = None

    document.add_heading(current_sheet, level=2)

    for i in range (first_line, ws.max_row):
        print("#####################################")
        cid = ws[i][0].value
        module = ws[i][1].value
        function = ws[i][2].value
        case = ws[i][3].value
        level = ws[i][4].value
        steps = ws[i][5].value
        presult = ws[i][6].value
        result = ws[i][7].value
        beizhu = ws[i][8].value

        if module == None:
            module = current_module
        else:
            document.add_heading(module, level=3)
            current_module = module
        
        if function == None:
            function = current_function
        else:
            document.add_heading(function, level=4)
            current_function = function
        
        records = [module, function, case, ' ', steps, presult, result, beizhu]
        # records.append({'module':str(module), 'function':str(function), 'case':str(case), 'top':'',
        # 'steps':RichText(str(steps)), 'presult':RichText(str(presult)), 'result':str(result), 'beizhu':str(beizhu)})
        # print(records)
        document.add_heading(case, level=5)
        write_doc_table(document, records)

out_dir = 'output/'
out_filename = out_dir + 'RG2018测试报告.docx'
document.save(out_filename)

# doc = DocxTemplate("templates/template2.docx")
# context = {
#     'contents':records,
# }
# doc.render(context)
# doc.save("output/generated_doc5.docx")