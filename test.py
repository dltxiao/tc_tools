from docx import Document
from docx.shared import Inches

doc1 = Document()

doc1.add_heading('Title1', level=1)
doc1.add_heading('Title2', level=2)
doc1.add_heading('Title3', level=3)

doc1.add_page_break()
table = doc1.add_table(rows=2, cols=2)
cell = table.cell(0,1)
cell.text = 'parrot, possibly dead'
row = table.rows[1]
row.cells[0].text = 'Foo bar to you.'
row.cells[1].text = 'And a hearty foo bar to you too sir!'

for row in table.rows:
    for cell in row.cells:
        print(cell.text)

row_count = len(table.rows)
col_count = len(table.columns)

row = table.add_row()

table.style = 'LightShading-Accent1'

doc1.add_heading('head1', level=3)

#####
table2 = doc1.add_table(rows=8, cols=0, style='Table Grid')
col1 = table2.add_column(Inches(1.3))
col1.cells[0].text = 'test'
col2 = table2.add_column(Inches(4.7))
col2.cells[1].text = '123'


doc1.save('output/doc1.docx')